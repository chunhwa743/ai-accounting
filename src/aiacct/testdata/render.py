"""Renders the synthetic dataset into the file formats a firm actually receives.

Real clients hand over a mix: a clean PDF from one bank, a phone photo of a
receipt, a scan someone made on a bad copier, a CSV export. Generating all of
them from one definition keeps the ground truth in a single place while still
exercising every ingestion path.

Everything is seeded, so regeneration is byte-identical apart from the JPEG
noise, which uses its own seeded generator.
"""

from __future__ import annotations

import csv
import json
import random
import textwrap
from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFilter
from reportlab import rl_config
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .models import SupportingDoc, Txn, period_dates, running_balances

PAGE_ROWS = 18  # rows per statement page, so multi-page handling is exercised

# Every file this module writes is a document by appearance - a statement, a tax
# invoice, a till receipt. The notice travels with the file, so a copy that has
# been separated from this repository still says what it is. Any real trading
# name appearing in a description is a nominative reference and nothing more:
# no affiliation, endorsement or actual transaction is represented.
SPECIMEN_NOTICE = (
    "SPECIMEN - synthetic data generated for software testing. "
    "This is not a genuine financial record. The account holder, the issuing "
    "institution and all amounts are fictional, and no real party is represented."
)

# reportlab stamps a creation timestamp and a random document id into every
# PDF, which would make regeneration produce different bytes each run. This
# fixes both, so "regenerate and diff" is a meaningful check.
rl_config.invariant = 1


def _fmt(value: Decimal | None) -> str:
    return "" if value is None else f"{value:,.2f}"


# ---------------------------------------------------------------- statements


def render_statement_pdf(
    path: Path,
    *,
    bank: str,
    client_name: str,
    account_number: str,
    period: str,
    txns: list[Txn],
    opening: Decimal,
) -> int:
    """A clean, text-layer PDF in the shape a Singapore bank issues."""
    start, end = period_dates(period)
    balances = running_balances(opening, txns)
    styles = getSampleStyleSheet()
    header = ParagraphStyle("hdr", parent=styles["Normal"], fontSize=8, leading=11)

    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
        title=f"{bank} Statement {period}",
    )

    flow = [
        Paragraph(f"<b><font size=13>{bank}</font></b>", styles["Normal"]),
        Paragraph("STATEMENT OF ACCOUNT", header),
        Spacer(1, 6 * mm),
        Paragraph(
            f"<b>{client_name}</b><br/>"
            f"Account Number: {account_number}<br/>"
            f"Statement Period: {start:%d %b %Y} to {end:%d %b %Y}<br/>"
            f"Currency: SGD",
            header,
        ),
        Spacer(1, 5 * mm),
    ]

    rows = [["Date", "Description", "Reference", "Withdrawal", "Deposit", "Balance"]]
    rows.append([
        f"{start:%d/%m/%Y}", "BALANCE BROUGHT FORWARD", "", "", "", _fmt(opening)
    ])
    for txn, balance in zip(txns, balances):
        rows.append([
            f"{date(start.year, start.month, txn.day):%d/%m/%Y}",
            txn.description,
            txn.reference or "",
            _fmt(txn.money_out),
            _fmt(txn.money_in),
            _fmt(balance),
        ])
    rows.append([
        f"{end:%d/%m/%Y}", "BALANCE CARRIED FORWARD", "", "", "",
        _fmt(balances[-1] if balances else opening),
    ])

    table = Table(
        rows,
        colWidths=[18 * mm, 72 * mm, 20 * mm, 22 * mm, 21 * mm, 21 * mm],
        repeatRows=1,
    )
    table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 7.5),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 7.5),
        ("LINEBELOW", (0, 0), (-1, 0), 0.75, colors.black),
        ("LINEABOVE", (0, -1), (-1, -1), 0.5, colors.black),
        ("ALIGN", (3, 1), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.Color(0.96, 0.96, 0.96)]),
    ]))
    flow.append(table)
    flow.append(Spacer(1, 4 * mm))
    flow.append(Paragraph(
        f"Total transactions: {len(txns)}<br/>"
        "This statement is computer generated. Please report any discrepancy "
        "within 14 days.",
        header,
    ))
    flow.append(Spacer(1, 4 * mm))
    flow.append(Paragraph(f"<b>{SPECIMEN_NOTICE}</b>", header))

    doc.build(flow)
    return max(1, (len(txns) // PAGE_ROWS) + 1)


def degrade_to_scan(source_pdf: Path, target_pdf: Path, seed: int = 7) -> None:
    """Turn a clean PDF into something that looks like it went through a copier.

    Skew, blur, speckle and JPEG artefacts, so the text layer disappears and
    the file has to go down the vision path rather than being read directly.
    """
    import pypdfium2

    rng = random.Random(seed)
    pdf = pypdfium2.PdfDocument(str(source_pdf))
    pages: list[Image.Image] = []

    for index in range(len(pdf)):
        rendered = pdf[index].render(scale=2.0).to_pil().convert("L")

        rendered = rendered.rotate(
            rng.uniform(-0.7, 0.7), resample=Image.BICUBIC, expand=False, fillcolor=245
        )
        rendered = rendered.filter(ImageFilter.GaussianBlur(radius=0.6))

        # Speckle and a faint vertical band, the way a dirty scanner roller looks.
        draw = ImageDraw.Draw(rendered)
        width, height = rendered.size
        for _ in range(int(width * height * 0.0006)):
            x, y = rng.randrange(width), rng.randrange(height)
            draw.point((x, y), fill=rng.randint(90, 180))
        band = rng.randrange(width)
        draw.line([(band, 0), (band, height)], fill=222, width=2)

        buffer = BytesIO()
        rendered.convert("RGB").save(buffer, format="JPEG", quality=42)
        buffer.seek(0)
        pages.append(Image.open(buffer).convert("RGB"))

    pdf.close()
    pages[0].save(target_pdf, save_all=True, append_images=pages[1:], resolution=150.0)


def render_statement_csv(
    path: Path, *, period: str, txns: list[Txn], opening: Decimal, with_balances: bool
) -> None:
    """A bank's CSV export.

    ``with_balances=False`` produces a file with no opening or closing figure,
    which makes the reconciliation check impossible to run - the case that has
    to be recorded as "unverifiable" rather than quietly passed.
    """
    start, _ = period_dates(period)
    balances = running_balances(opening, txns)

    with open(path, "w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        header = ["Transaction Date", "Description", "Reference", "Debit Amount", "Credit Amount"]
        if with_balances:
            header.append("Balance")
        writer.writerow(header)

        for txn, balance in zip(txns, balances):
            row = [
                f"{date(start.year, start.month, txn.day):%d/%m/%Y}",
                txn.description,
                txn.reference or "",
                _fmt(txn.money_out),
                _fmt(txn.money_in),
            ]
            if with_balances:
                row.append(_fmt(balance))
            writer.writerow(row)

        # A trailing row, not a leading one: the reader takes row 0 as the
        # header, and already drops later rows that carry no parseable date.
        writer.writerow([])
        writer.writerow([f"# {SPECIMEN_NOTICE}"])


# ---------------------------------------------------------------- documents


def render_invoice_pdf(path: Path, doc: SupportingDoc, buyer: str) -> None:
    styles = getSampleStyleSheet()
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8.5, leading=12)

    template = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm,
        title=f"{doc.vendor_name} {doc.doc_number}",
    )

    net = doc.total - doc.tax
    rows = [["Description", "Amount (SGD)"]]
    for item in doc.line_items or [doc.summary]:
        rows.append([item, ""])
    rows.append(["Subtotal", _fmt(net)])
    rows.append(["GST @ 9%", _fmt(doc.tax)])
    rows.append(["Total", _fmt(doc.total)])

    table = Table(rows, colWidths=[120 * mm, 40 * mm])
    table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("FONT", (0, -1), (-1, -1), "Helvetica-Bold", 9),
        ("LINEBELOW", (0, 0), (-1, 0), 0.75, colors.black),
        ("LINEABOVE", (0, -3), (-1, -3), 0.5, colors.black),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
    ]))

    template.build([
        Paragraph(f"<b><font size=13>{doc.vendor_name}</font></b>", styles["Normal"]),
        Paragraph("TAX INVOICE", small),
        Spacer(1, 5 * mm),
        Paragraph(
            f"GST Reg No: 20{doc.doc_number[-6:]}X<br/>"
            f"Invoice No: <b>{doc.doc_number}</b><br/>"
            f"Date: {doc.doc_date:%d %B %Y}<br/><br/>"
            f"Bill To:<br/><b>{buyer}</b>",
            small,
        ),
        Spacer(1, 6 * mm),
        table,
        Spacer(1, 6 * mm),
        Paragraph("Payment terms: 30 days. Please quote the invoice number.", small),
        Spacer(1, 5 * mm),
        Paragraph(f"<b>{SPECIMEN_NOTICE}</b>", small),
    ])


def render_receipt_jpg(path: Path, doc: SupportingDoc, seed: int = 11) -> None:
    """A phone photo of a till receipt - narrow, mono, slightly skewed."""
    rng = random.Random(seed)
    width, height = 520, 760
    image = Image.new("L", (width, height), color=250)
    draw = ImageDraw.Draw(image)

    net = doc.total - doc.tax
    lines = [
        doc.vendor_name.upper(), "", f"RECEIPT {doc.doc_number}",
        f"{doc.doc_date:%d/%m/%Y}", "-" * 34,
    ]
    for item in doc.line_items or [doc.summary]:
        lines.append(item[:34])
    lines += [
        "-" * 34,
        f"SUBTOTAL{_fmt(net):>26}",
        f"GST 9%{_fmt(doc.tax):>28}",
        f"TOTAL{_fmt(doc.total):>29}",
        "-" * 34, "", "NETS  ****4471", "THANK YOU",
        "", "-" * 34,
        *textwrap.wrap(SPECIMEN_NOTICE, width=34),
    ]

    y = 40
    for line in lines:
        draw.text((40, y), line, fill=35)
        y += 26

    image = image.rotate(rng.uniform(-2.0, 2.0), resample=Image.BICUBIC, fillcolor=250)
    image = image.filter(ImageFilter.GaussianBlur(radius=0.4))
    for _ in range(1400):
        draw = ImageDraw.Draw(image)
        draw.point((rng.randrange(width), rng.randrange(height)), fill=rng.randint(120, 200))

    image.convert("RGB").save(path, format="JPEG", quality=70)


def render_payroll_docx(path: Path, doc: SupportingDoc, employees: int = 6) -> None:
    document = DocxDocument()
    document.add_heading(doc.vendor_name, level=1)
    document.add_paragraph(f"Payroll Summary - {doc.doc_date:%B %Y}")
    document.add_paragraph(f"Reference: {doc.doc_number}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for cell, label in zip(table.rows[0].cells, ["Employee", "Gross", "CPF (Employee)", "Net"]):
        cell.text = label

    gross_each = doc.total / employees
    for index in range(employees):
        cells = table.add_row().cells
        cells[0].text = f"Employee {index + 1:02d}"
        cells[1].text = _fmt(gross_each)
        cells[2].text = _fmt(gross_each * Decimal("0.20"))
        cells[3].text = _fmt(gross_each * Decimal("0.80"))

    document.add_paragraph("")
    document.add_paragraph(f"Total net pay: SGD {_fmt(doc.total)}")
    document.add_paragraph(
        "Employer CPF contributions are remitted separately to the CPF Board."
    )
    document.add_paragraph("")
    document.add_paragraph(SPECIMEN_NOTICE).runs[0].bold = True
    document.save(path)


# ---------------------------------------------------------------- sidecars


def statement_sidecar(
    *,
    bank: str,
    account_number: str,
    account_holder: str,
    period: str,
    txns: list[Txn],
    opening: Decimal,
    with_balances: bool,
    unclear_header: dict[str, str] | None = None,
) -> dict:
    """The transcription the offline provider replays.

    This is what is *printed* on the statement, not the answer key - the
    categorisation still has to be worked out. Without it the stub could not
    read a PDF at all and the offline demo would be impossible.
    """
    start, _ = period_dates(period)
    balances = running_balances(opening, txns)

    transactions = []
    for index, (txn, balance) in enumerate(zip(txns, balances), start=1):
        transactions.append({
            "line_no": index,
            "txn_date": date(start.year, start.month, txn.day).isoformat(),
            "raw_description": txn.description,
            "bank_reference": txn.reference,
            "money_in": float(txn.money_in) if txn.money_in is not None else None,
            "money_out": float(txn.money_out) if txn.money_out is not None else None,
            "balance_after": float(balance) if with_balances else None,
            "page": (index - 1) // PAGE_ROWS + 1,
            "uncertain_fields": [],
        })

    return {
        "bank_name": bank,
        "account_number": account_number,
        "account_holder": account_holder,
        "period_start": start.isoformat(),
        "period_end": period_dates(period)[1].isoformat(),
        "opening_balance": float(opening) if with_balances else None,
        "closing_balance": float(balances[-1]) if (with_balances and balances) else None,
        "stated_transaction_count": len(txns),
        "transactions": transactions,
        "uncertain_fields": [
            {"field": name, "legibility": value}
            for name, value in (unclear_header or {}).items()
        ],
    }


def document_sidecar(doc: SupportingDoc, unclear: dict[str, str] | None = None) -> dict:
    return {
        "vendor_name": doc.vendor_name,
        "doc_number": doc.doc_number,
        "doc_date": doc.doc_date.isoformat(),
        "total_amount": float(doc.total),
        "tax_amount": float(doc.tax),
        "summary": doc.summary,
        "uncertain_fields": [
            {"field": name, "legibility": value} for name, value in (unclear or {}).items()
        ],
    }


def write_sidecar(target: Path, payload: dict) -> None:
    sidecar = target.with_suffix(target.suffix + ".extract.json")
    sidecar.write_text(json.dumps(payload, indent=2), encoding="utf-8")
