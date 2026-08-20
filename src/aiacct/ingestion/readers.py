"""Local readers for the formats that need no model at all.

A CSV is already structured; a text-layer PDF already contains its own words.
Sending either to a model would be paying for something arithmetic and
string handling do perfectly.
"""

from __future__ import annotations

import csv
import re
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path

import logging

from ..models import ColumnMapping, ExtractedTransaction, StatementExtraction

log = logging.getLogger(__name__)

# Bank CSV exports do not agree on column names, so headers are matched by
# meaning rather than by an exact string.
COLUMN_ALIASES: dict[str, tuple[str, ...]] = {
    "date": ("transaction date", "txn date", "date", "posting date", "value date"),
    "description": ("description", "particulars", "transaction details", "narrative", "details"),
    "reference": ("reference", "ref", "cheque", "transaction ref"),
    "debit": ("debit amount", "withdrawal", "debit", "money out", "paid out"),
    "credit": ("credit amount", "deposit", "credit", "money in", "paid in"),
    "balance": ("balance", "running balance", "ledger balance"),
}

# Unambiguous formats first. Slash-separated dates are handled separately,
# because "03/05/2026" is 3 May or 5 March depending on where the file came
# from and guessing produces a wrong answer that nothing downstream can see.
DATE_FORMATS = ("%Y-%m-%d", "%d %b %Y", "%d %B %Y", "%b %d %Y", "%Y/%m/%d")
SLASH_DATE = re.compile(r"^(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})$")


def parse_amount(raw: str | None) -> Decimal | None:
    if raw is None:
        return None
    text = str(raw).strip().replace(",", "").replace("$", "").replace("SGD", "").strip()
    if not text or text in {"-", "--"}:
        return None
    negative = text.startswith("(") and text.endswith(")")
    if negative:
        text = text[1:-1]
    try:
        value = Decimal(text)
    except InvalidOperation:
        return None
    return -value if negative else value


def parse_date(raw: str | None, day_first: bool = True) -> date | None:
    """Parse one date. ``day_first`` resolves the ambiguous slash forms.

    Callers that have a whole column available should use
    :func:`detect_day_first` rather than accepting the default, because
    guessing wrong shifts a transaction into the wrong month silently - the
    balance checks cannot see it, since dates take no part in the arithmetic.
    """
    if not raw:
        return None
    text = str(raw).strip()

    for fmt in DATE_FORMATS:
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue

    match = SLASH_DATE.match(text)
    if match:
        first, second, year = (int(g) for g in match.groups())
        if year < 100:
            year += 2000
        day, month = (first, second) if day_first else (second, first)
        try:
            return date(year, month, day)
        except ValueError:
            # The assumed order is impossible - month 25 - so it must be the
            # other way round.
            try:
                return date(year, day, month)
            except ValueError:
                return None

    try:
        return date.fromisoformat(text[:10])
    except ValueError:
        return None


def detect_day_first(values: list[str]) -> tuple[bool, bool]:
    """Infer date order from a whole column, and say whether it was inferable.

    Returns ``(day_first, certain)``. Any value with a first component above 12
    settles it for the file; if every value is ambiguous the caller is told so
    rather than being handed a confident guess.
    """
    saw_day_first = saw_month_first = False
    for value in values:
        match = SLASH_DATE.match(str(value).strip())
        if not match:
            continue
        first, second, _ = (int(g) for g in match.groups())
        if first > 12:
            saw_day_first = True
        elif second > 12:
            saw_month_first = True

    if saw_day_first and not saw_month_first:
        return True, True
    if saw_month_first and not saw_day_first:
        return False, True
    # Either nothing settled it, or the column contains both - which means the
    # file is inconsistent and worth flagging.
    return True, False


def map_columns(header: list[str]) -> dict[str, int]:
    """Match header cells to meanings, longest alias first.

    "Transaction Date" must not be captured by the shorter "date" alias when
    both are present.
    """
    normalised = [h.strip().lower() for h in header]
    mapping: dict[str, int] = {}
    for field, aliases in COLUMN_ALIASES.items():
        for alias in sorted(aliases, key=len, reverse=True):
            for index, cell in enumerate(normalised):
                if index in mapping.values():
                    continue
                if cell == alias or alias in cell:
                    mapping[field] = index
                    break
            if field in mapping:
                break
    return mapping


COLUMN_PROMPT = """This is the top of a bank's CSV or spreadsheet export. Identify which column
holds each field, by its zero-based position.

Return null for anything the file does not have. Many exports use a single
signed amount column rather than separate debit and credit columns; if so, map
it to `debit` and say so in `notes`.

<rows>
{rows}
</rows>
"""


def _map_columns_with_model(rows: list[list[str]], llm, filename: str) -> dict[str, int] | None:
    """Ask the model to map an unfamiliar header.

    The alias table covers the layouts we have seen; this covers the ones we
    have not, which is every bank whose export nobody has looked at yet. One
    cheap call on a handful of rows beats maintaining a list that fails loudly
    on the first unfamiliar file - or worse, quietly maps the wrong column.
    """
    if llm is None:
        return None

    sample = chr(10).join(", ".join(row) for row in rows[:4])
    try:
        result = llm.parse(
            prompt=COLUMN_PROMPT.format(rows=sample),
            schema=ColumnMapping,
            effort="low",
        )
    except Exception as exc:  # noqa: BLE001
        log.warning("%s: column mapping call failed (%s)", filename, exc)
        return None

    mapping = {
        field: index
        for field, index in (
            ("date", result.parsed.date_column),
            ("description", result.parsed.description_column),
            ("reference", result.parsed.reference_column),
            ("debit", result.parsed.debit_column),
            ("credit", result.parsed.credit_column),
            ("balance", result.parsed.balance_column),
        )
        if index is not None and 0 <= index < len(rows[0])
    }
    log.info("%s: header mapped by the model -> %s", filename, mapping)
    return mapping or None


def read_tabular_statement(path: Path, llm=None) -> StatementExtraction:
    """Parse a CSV or XLSX bank export.

    The header alias table handles the common cases for nothing. When it does
    not recognise a layout - and bank exports vary more than any fixed list can
    anticipate - the header and a few sample rows go to the model instead of
    the parse failing outright.
    """
    rows = _load_rows(path)
    if not rows:
        raise ValueError(f"{path.name} is empty")

    mapping = map_columns(rows[0])
    missing = {"date", "description"} - set(mapping)
    if missing:
        mapping = _map_columns_with_model(rows, llm, path.name) or mapping
        missing = {"date", "description"} - set(mapping)
    if missing:
        raise ValueError(
            f"{path.name}: could not identify the {sorted(missing)} column(s) "
            f"in header {rows[0]}"
        )

    date_index = mapping["date"]
    day_first, certain = detect_day_first(
        [row[date_index] for row in rows[1:] if date_index < len(row)]
    )
    if not certain:
        log.warning(
            "%s: every date is ambiguous (no day above 12), assuming day-first. "
            "A month-first export would be silently misread.", path.name,
        )

    transactions: list[ExtractedTransaction] = []
    for line_no, row in enumerate(rows[1:], start=1):
        if not any(cell.strip() for cell in row):
            continue

        def cell(field: str) -> str | None:
            index = mapping.get(field)
            return row[index] if index is not None and index < len(row) else None

        txn_date = parse_date(cell("date"), day_first=day_first)
        if txn_date is None:
            continue  # subtotal or footer row

        debit = parse_amount(cell("debit"))
        credit = parse_amount(cell("credit"))
        balance = parse_amount(cell("balance"))

        transactions.append(
            ExtractedTransaction(
                line_no=line_no,
                txn_date=txn_date.isoformat(),
                raw_description=(cell("description") or "").strip(),
                bank_reference=(cell("reference") or "").strip() or None,
                money_in=float(credit) if credit else None,
                money_out=float(debit) if debit else None,
                balance_after=float(balance) if balance is not None else None,
                page=1,
                # A parsed file has no legibility problem: either the column was
                # there or it was not.
                uncertain_fields=[],
            )
        )

    opening = None
    closing = None
    if transactions and transactions[0].balance_after is not None:
        first = transactions[0]
        movement = (first.money_in or 0) - (first.money_out or 0)
        opening = round(first.balance_after - movement, 2)
        closing = transactions[-1].balance_after

    return StatementExtraction(
        bank_name=None,
        account_number=None,
        account_holder=None,
        period_start=transactions[0].txn_date if transactions else None,
        period_end=transactions[-1].txn_date if transactions else None,
        opening_balance=opening,
        closing_balance=closing,
        stated_transaction_count=None,
        transactions=transactions,
        uncertain_fields=[],
    )


def _load_rows(path: Path) -> list[list[str]]:
    if path.suffix.lower() in {".xlsx", ".xls"}:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        sheet = workbook.active
        rows = [
            ["" if cell is None else str(cell) for cell in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        workbook.close()
        return rows

    with open(path, newline="", encoding="utf-8-sig") as handle:
        sample = handle.read(4096)
        handle.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        return [row for row in csv.reader(handle, dialect)]


def read_pdf_text(path: Path, max_pages: int | None = None) -> str:
    """Extract the text layer, keeping page markers.

    The markers matter: when a check fails we need to say which page to look
    at, and the model needs to report which page each row came from.
    """
    import pdfplumber

    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        pages = pdf.pages[:max_pages] if max_pages else pdf.pages
        for index, page in enumerate(pages, start=1):
            chunks.append(f"--- PAGE {index} ---")
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def read_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_first_page_text(path: Path, kind: str) -> str:
    """Just enough to classify a document, without reading the whole thing."""
    if kind == "pdf":
        return read_pdf_text(path, max_pages=1)[:4000]
    if kind == "docx":
        return read_docx_text(path)[:4000]
    if kind == "tabular":
        rows = _load_rows(path)[:20]
        return "\n".join(", ".join(row) for row in rows)[:4000]
    return path.read_text(encoding="utf-8", errors="replace")[:4000]


_WHITESPACE = re.compile(r"\s+")


def tidy(text: str | None) -> str | None:
    return _WHITESPACE.sub(" ", text).strip() if text else None
